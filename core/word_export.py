"""
core/word_export.py
===================
Génération du rapport Word structuré par thème, en Python pur avec python-docx.
Aucune dépendance Node.js ou npm — fonctionne sur Streamlit Cloud sans configuration.
"""

import io
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime


# ── Palette couleurs par thème ────────────────────────────────────────────────
THEME_META = {
    "avifaune":      {"label": "🦅 Avifaune",       "hex": "1D4ED8", "light": "DBEAFE"},
    "chiropteres":   {"label": "🦇 Chiroptères",     "hex": "7C3AED", "light": "EDE9FE"},
    "zones_humides": {"label": "💧 Zones humides",   "hex": "0F766E", "light": "CCFBF1"},
    "paysage":       {"label": "🌄 Paysage",         "hex": "C2410C", "light": "FFEDD5"},
}
THEME_ORDER = ["avifaune", "chiropteres", "zones_humides", "paysage"]
VERT        = "166534"


def _rgb(hex_str: str) -> RGBColor:
    r, g, b = int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16)
    return RGBColor(r, g, b)


def _set_cell_bg(cell, hex_color: str):
    """Colorie le fond d'une cellule de tableau."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, left_color: str, left_size: int = 12):
    """Ajoute une bordure gauche colorée et épaisse à une cellule."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "2")
        el.set(qn("w:color"), "D1D5DB")
        tcBorders.append(el)
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    str(left_size))
    left.set(qn("w:color"), left_color)
    tcBorders.append(left)
    tcPr.append(tcBorders)


def _score_stars(score: int) -> str:
    return "●" * min(score // 4 + 1, 5)


def _add_horizontal_rule(doc: Document, color: str = "D1D5DB"):
    """Ajoute un filet horizontal via bordure basse de paragraphe."""
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:color"), color)
    bot.set(qn("w:space"), "1")
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.space_before = Pt(0)
    return p


# ── Page de garde ─────────────────────────────────────────────────────────────
def _make_cover(doc: Document, meta: dict):
    for _ in range(4):
        doc.add_paragraph()

    # Titre
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run("Analyse environnementale")
    run.font.size  = Pt(26)
    run.font.bold  = True
    run.font.color.rgb = _rgb(VERT)
    h.paragraph_format.space_after = Pt(6)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Arrêtés préfectoraux éoliens")
    r2.font.size  = Pt(18)
    r2.font.color.rgb = _rgb("374151")
    sub.paragraph_format.space_after = Pt(18)

    _add_horizontal_rule(doc, VERT)

    # Métadonnées
    date_str = datetime.now().strftime("%d %B %Y")
    for label, value in [
        ("Région",    meta.get("region") or "Non spécifiée"),
        ("Documents", str(meta.get("nb_docs",    0))),
        ("Extraits",  str(meta.get("nb_passages", 0))),
        ("Généré le", date_str),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        rl = p.add_run(f"{label} : ")
        rl.font.bold  = True
        rl.font.color.rgb = _rgb("6B7280")
        rl.font.size  = Pt(11)
        rv = p.add_run(value)
        rv.font.color.rgb = _rgb("111827")
        rv.font.size  = Pt(11)

    doc.add_page_break()


# ── Tableau de synthèse ───────────────────────────────────────────────────────
def _make_summary(doc: Document, results: list):
    h = doc.add_heading("Synthèse par thématique", level=1)
    h.runs[0].font.color.rgb = _rgb(VERT)
    h.paragraph_format.space_after = Pt(10)

    # Calculer les stats par thème
    stats = {}
    for doc_result in results:
        for p in (doc_result.get("passages") or []):
            for th, score in (p.get("themes") or {}).items():
                if th not in stats:
                    stats[th] = {"count": 0, "docs": set(), "max_score": 0}
                stats[th]["count"]    += 1
                stats[th]["docs"].add(doc_result.get("text") or doc_result.get("filename") or "")
                stats[th]["max_score"] = max(stats[th]["max_score"], score)

    # Tableau 4 colonnes
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Largeurs colonnes (en twips, 1440 twips = 1 pouce, A4 contenu ≈ 9026 twips)
    widths = [2800, 1600, 2200, 2200]
    for i, cell in enumerate(table.rows[0].cells):
        cell.width = Twips(widths[i])

    # En-tête
    for i, (cell, header) in enumerate(zip(table.rows[0].cells,
                                           ["Thématique", "Passages", "Documents", "Score max"])):
        _set_cell_bg(cell, VERT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size  = Pt(10)

    # Lignes de données
    for th_key in THEME_ORDER:
        tm = THEME_META.get(th_key, {})
        s  = stats.get(th_key, {"count": 0, "docs": set(), "max_score": 0})
        row = table.add_row()
        for i, cell in enumerate(row.cells):
            cell.width = Twips(widths[i])

        values = [
            tm.get("label", th_key),
            str(s["count"]),
            str(len(s["docs"])),
            f"{s['max_score']} {_score_stars(s['max_score'])}" if s["max_score"] else "—",
        ]
        bg_colors = [tm.get("light", "FFFFFF"), "FFFFFF", "FFFFFF", "FFFFFF"]

        for i, (cell, val) in enumerate(zip(row.cells, values)):
            _set_cell_bg(cell, bg_colors[i])
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(10)
            if i == 0:
                run.font.bold = True
                run.font.color.rgb = _rgb(tm.get("hex", "374151"))

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    doc.add_page_break()


# ── Section d'un thème ────────────────────────────────────────────────────────
def _make_theme_section(doc: Document, theme_key: str, results: list):
    tm    = THEME_META.get(theme_key, {"label": theme_key, "hex": "374151", "light": "F9FAFB"})
    color = tm["hex"]

    # Titre de section
    h = doc.add_heading(tm["label"], level=1)
    h.runs[0].font.color.rgb = _rgb(color)
    h.paragraph_format.space_after = Pt(2)
    _add_horizontal_rule(doc, color)

    # Collecter et trier les passages
    all_passages = []
    for doc_result in results:
        for p in (doc_result.get("passages") or []):
            if theme_key in (p.get("themes") or {}):
                all_passages.append({
                    **p,
                    "_score":      p["themes"][theme_key],
                    "_doc":        doc_result.get("text") or doc_result.get("filename") or "Document inconnu",
                    "_region":     doc_result.get("region", ""),
                    "_date":       (f"{doc_result['month']}/{doc_result['year']}"
                                   if doc_result.get("year") else ""),
                    "type_long":   doc_result.get("type_long",  ""),
                    "type_found":  doc_result.get("type_found", False),
                })
    all_passages.sort(key=lambda x: x["_score"], reverse=True)

    if not all_passages:
        p = doc.add_paragraph("Aucun extrait détecté pour ce thème.")
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = _rgb("6B7280")
        doc.add_page_break()
        return

    # Grouper par document
    by_doc: dict = {}
    for p in all_passages:
        key = p["_doc"]
        if key not in by_doc:
            by_doc[key] = {"meta": p, "passages": []}
        by_doc[key]["passages"].append(p)

    for doc_name, group in by_doc.items():
        # Sous-titre document
        h2 = doc.add_heading(doc_name[:100], level=2)
        h2.runs[0].font.color.rgb = _rgb("374151")
        h2.paragraph_format.space_before = Pt(12)
        h2.paragraph_format.space_after  = Pt(3)

        meta = group["meta"]
        info_parts = []
        if meta.get("type_long") and meta.get("type_found"):
            info_parts.append(meta["type_long"])
        if meta["_region"]:
            info_parts.append(meta["_region"])
        if meta["_date"]:
            info_parts.append(meta["_date"])
        if info_parts:
            info = doc.add_paragraph()
            r = info.add_run(" · ".join(info_parts))
            r.font.italic = True
            r.font.color.rgb = _rgb("6B7280")
            r.font.size = Pt(9)
            info.paragraph_format.space_after = Pt(6)

        for p in group["passages"]:
            # Ligne article + thèmes + score
            meta_p = doc.add_paragraph()
            meta_p.paragraph_format.space_before = Pt(6)
            meta_p.paragraph_format.space_after  = Pt(2)

            if p.get("article_ref"):
                art_run = meta_p.add_run(f"{p['article_ref']}  ")
                art_run.font.bold  = True
                art_run.font.color.rgb = _rgb(color)
                art_run.font.size  = Pt(9)

            all_themes_str = " · ".join(
                THEME_META.get(k, {}).get("label", k)
                for k in (p.get("themes") or {})
            )
            th_run = meta_p.add_run(all_themes_str)
            th_run.font.bold  = True
            th_run.font.color.rgb = _rgb(color)
            th_run.font.size  = Pt(9)

            sc_run = meta_p.add_run(f"   Score : {p['_score']} {_score_stars(p['_score'])}")
            sc_run.font.color.rgb = _rgb("6B7280")
            sc_run.font.size  = Pt(9)

            # Extrait dans un tableau 1×1 avec bordure gauche colorée
            tbl  = doc.add_table(rows=1, cols=1)
            tbl.style = "Table Grid"
            cell = tbl.rows[0].cells[0]
            cell.width = Twips(9000)
            _set_cell_bg(cell,    tm["light"])
            _set_cell_border(cell, color, left_size=18)

            text_content = " ".join((p.get("text") or "").split())
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_after  = Pt(0)
            cp.paragraph_format.left_indent  = Cm(0.3)
            cr = cp.add_run(text_content)
            cr.font.size  = Pt(10)
            cr.font.color.rgb = _rgb("111827")

            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(6)

    doc.add_page_break()


# ── Point d'entrée public ─────────────────────────────────────────────────────
def generate_word_report(results: list) -> bytes:
    """
    Génère un rapport Word structuré par thème.
    Retourne les bytes du fichier .docx.
    """
    doc = Document()

    # Marges 2 cm sur tous les côtés
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    # Style de base
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Styles de titres
    for level, (size, space_before, space_after) in enumerate(
        [(18, 12, 6), (13, 10, 4)], start=1
    ):
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name  = "Arial"
        h_style.font.size  = Pt(size)
        h_style.font.bold  = True
        h_style.font.color.rgb = _rgb("111827")
        h_style.paragraph_format.space_before = Pt(space_before)
        h_style.paragraph_format.space_after  = Pt(space_after)

    meta = {
        "region":      results[0].get("region", "") if results else "",
        "nb_docs":     len(results),
        "nb_passages": sum(len(r.get("passages") or []) for r in results),
    }

    _make_cover(doc, meta)
    _make_summary(doc, results)
    for th in THEME_ORDER:
        _make_theme_section(doc, th, results)

    # Sauvegarder en mémoire
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
